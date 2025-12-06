"""
Advanced Document Analyzer - Enterprise-grade document analysis features.

Provides comprehensive document analysis including:
- Document Metadata Extraction (author, creation date, company, software)
- Document Tampering/Integrity Detection
- Language Detection
- Document Classification
- Hidden Content Detection
- Embedded Object Detection
- Digital Signature Detection
- Redaction Detection
- Version/Revision History Analysis
- Enhanced Table/Form Extraction
- Document Forensics
- Cross-Reference Detection
"""

import os
import re
import logging
import hashlib
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from io import BytesIO

try:
    from utils.centralized_logger import get_scanner_logger
    logger = get_scanner_logger("advanced_document_analyzer")
except ImportError:
    logger = logging.getLogger(__name__)

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import olefile
    OLE_AVAILABLE = True
except ImportError:
    OLE_AVAILABLE = False

try:
    from langdetect import detect, detect_langs
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


class AdvancedDocumentAnalyzer:
    """
    Enterprise-grade document analyzer with 12 advanced detection features.
    """
    
    def __init__(self, region: str = "Netherlands"):
        self.region = region
        self.analysis_features = {
            'metadata_extraction': True,
            'tampering_detection': True,
            'language_detection': True,
            'classification': True,
            'hidden_content': True,
            'embedded_objects': True,
            'digital_signatures': True,
            'redaction_detection': True,
            'version_history': True,
            'table_extraction': True,
            'forensics': True,
            'cross_references': True
        }
        
        self.document_categories = [
            'contract', 'invoice', 'medical', 'hr', 'legal', 
            'financial', 'personal', 'technical', 'marketing', 'general'
        ]
        
        self.classification_keywords = {
            'contract': ['agreement', 'contract', 'terms', 'conditions', 'parties', 'hereby', 'obligations', 'termination', 'effective date', 'signature', 'witness'],
            'invoice': ['invoice', 'bill', 'payment', 'amount due', 'total', 'tax', 'subtotal', 'vat', 'btw', 'due date', 'invoice number', 'qty', 'quantity'],
            'medical': ['patient', 'diagnosis', 'treatment', 'medication', 'prescription', 'doctor', 'hospital', 'medical', 'health', 'symptoms', 'examination'],
            'hr': ['employee', 'salary', 'employment', 'hire', 'termination', 'benefits', 'vacation', 'leave', 'performance', 'review', 'personnel'],
            'legal': ['court', 'judge', 'plaintiff', 'defendant', 'attorney', 'lawyer', 'verdict', 'lawsuit', 'legal', 'jurisdiction', 'statute'],
            'financial': ['balance', 'assets', 'liabilities', 'revenue', 'profit', 'loss', 'quarterly', 'annual', 'fiscal', 'dividend', 'investment'],
            'personal': ['resume', 'cv', 'curriculum vitae', 'passport', 'identity', 'birth certificate', 'driver license', 'personal details'],
            'technical': ['specification', 'architecture', 'api', 'documentation', 'implementation', 'system', 'database', 'code', 'algorithm'],
            'marketing': ['campaign', 'marketing', 'brand', 'customer', 'advertisement', 'promotion', 'sales', 'target audience', 'roi']
        }
    
    def analyze_document(self, file_path: str, text_content: str = None) -> Dict[str, Any]:
        """
        Perform comprehensive document analysis.
        
        Args:
            file_path: Path to the document file
            text_content: Pre-extracted text content (optional)
            
        Returns:
            Dictionary containing all analysis results
        """
        if not os.path.exists(file_path):
            return {'error': 'File not found', 'file_path': file_path}
        
        _, ext = os.path.splitext(file_path.lower())
        file_name = os.path.basename(file_path)
        
        analysis_result = {
            'file_name': file_name,
            'file_path': file_path,
            'file_extension': ext,
            'analysis_timestamp': datetime.now().isoformat(),
            'metadata': {},
            'integrity': {},
            'language': {},
            'classification': {},
            'hidden_content': {},
            'embedded_objects': {},
            'digital_signatures': {},
            'redaction': {},
            'version_history': {},
            'table_data': {},
            'forensics': {},
            'cross_references': {},
            'findings': [],
            'risk_indicators': []
        }
        
        try:
            analysis_result['metadata'] = self._extract_metadata(file_path, ext)
            analysis_result['integrity'] = self._detect_tampering(file_path, ext, analysis_result['metadata'])
            
            if text_content:
                analysis_result['language'] = self._detect_language(text_content)
                analysis_result['classification'] = self._classify_document(text_content, file_name, ext)
                analysis_result['redaction'] = self._detect_redaction(text_content, file_path, ext)
                analysis_result['cross_references'] = self._detect_cross_references(text_content, file_path)
            
            analysis_result['hidden_content'] = self._detect_hidden_content(file_path, ext)
            analysis_result['embedded_objects'] = self._detect_embedded_objects(file_path, ext)
            analysis_result['digital_signatures'] = self._detect_digital_signatures(file_path, ext)
            analysis_result['version_history'] = self._analyze_version_history(file_path, ext)
            analysis_result['table_data'] = self._extract_table_data(file_path, ext)
            analysis_result['forensics'] = self._perform_forensic_analysis(file_path, ext, analysis_result['metadata'])
            
            analysis_result['findings'] = self._generate_findings(analysis_result)
            analysis_result['risk_indicators'] = self._calculate_risk_indicators(analysis_result)
            
        except Exception as e:
            logger.error(f"Error analyzing document {file_path}: {str(e)}")
            analysis_result['error'] = str(e)
        
        return analysis_result
    
    def _extract_metadata(self, file_path: str, ext: str) -> Dict[str, Any]:
        """
        Extract comprehensive metadata from document.
        """
        metadata = {
            'author': None,
            'creator': None,
            'producer': None,
            'creation_date': None,
            'modification_date': None,
            'last_modified_by': None,
            'company': None,
            'title': None,
            'subject': None,
            'keywords': [],
            'application': None,
            'version': None,
            'page_count': None,
            'word_count': None,
            'character_count': None,
            'file_size': os.path.getsize(file_path),
            'file_created': datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
            'file_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            'pii_in_metadata': []
        }
        
        try:
            if ext == '.pdf' and PDF_AVAILABLE:
                metadata.update(self._extract_pdf_metadata(file_path))
            elif ext in ['.docx', '.doc'] and DOCX_AVAILABLE:
                metadata.update(self._extract_docx_metadata(file_path))
            elif ext in ['.xlsx', '.xls'] and EXCEL_AVAILABLE:
                metadata.update(self._extract_excel_metadata(file_path))
            elif ext in ['.doc', '.xls', '.ppt'] and OLE_AVAILABLE:
                metadata.update(self._extract_ole_metadata(file_path))
            
            metadata['pii_in_metadata'] = self._detect_pii_in_metadata(metadata)
            
        except Exception as e:
            logger.warning(f"Metadata extraction failed for {file_path}: {str(e)}")
            metadata['extraction_error'] = str(e)
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        metadata = {}
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                info = reader.metadata
                
                if info:
                    metadata['author'] = info.get('/Author', None)
                    metadata['creator'] = info.get('/Creator', None)
                    metadata['producer'] = info.get('/Producer', None)
                    metadata['title'] = info.get('/Title', None)
                    metadata['subject'] = info.get('/Subject', None)
                    
                    if info.get('/CreationDate'):
                        metadata['creation_date'] = self._parse_pdf_date(info.get('/CreationDate'))
                    if info.get('/ModDate'):
                        metadata['modification_date'] = self._parse_pdf_date(info.get('/ModDate'))
                    
                    if info.get('/Keywords'):
                        keywords_str = info.get('/Keywords', '')
                        metadata['keywords'] = [k.strip() for k in keywords_str.split(',') if k.strip()]
                
                metadata['page_count'] = len(reader.pages)
                metadata['is_encrypted'] = reader.is_encrypted
                
        except Exception as e:
            logger.warning(f"PDF metadata extraction error: {str(e)}")
        
        return metadata
    
    def _parse_pdf_date(self, date_str: str) -> Optional[str]:
        """Parse PDF date format (D:YYYYMMDDHHmmss)."""
        if not date_str:
            return None
        try:
            date_str = str(date_str)
            if date_str.startswith('D:'):
                date_str = date_str[2:]
            if len(date_str) >= 14:
                dt = datetime.strptime(date_str[:14], '%Y%m%d%H%M%S')
                return dt.isoformat()
        except:
            pass
        return date_str
    
    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX files."""
        metadata = {}
        try:
            doc = docx.Document(file_path)
            props = doc.core_properties
            
            metadata['author'] = props.author
            metadata['last_modified_by'] = props.last_modified_by
            metadata['creation_date'] = props.created.isoformat() if props.created else None
            metadata['modification_date'] = props.modified.isoformat() if props.modified else None
            metadata['title'] = props.title
            metadata['subject'] = props.subject
            metadata['keywords'] = [k.strip() for k in (props.keywords or '').split(',') if k.strip()]
            metadata['company'] = getattr(props, 'company', None)
            metadata['version'] = props.version
            metadata['revision'] = props.revision
            metadata['category'] = props.category
            
            word_count = 0
            for para in doc.paragraphs:
                word_count += len(para.text.split())
            metadata['word_count'] = word_count
            metadata['page_count'] = len(doc.sections)
            
        except Exception as e:
            logger.warning(f"DOCX metadata extraction error: {str(e)}")
        
        return metadata
    
    def _extract_excel_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Excel files."""
        metadata = {}
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True)
            props = wb.properties
            
            metadata['author'] = props.creator
            metadata['last_modified_by'] = props.lastModifiedBy
            metadata['creation_date'] = props.created.isoformat() if props.created else None
            metadata['modification_date'] = props.modified.isoformat() if props.modified else None
            metadata['title'] = props.title
            metadata['subject'] = props.subject
            metadata['keywords'] = [k.strip() for k in (props.keywords or '').split(',') if k.strip()]
            metadata['company'] = props.company
            metadata['sheet_count'] = len(wb.sheetnames)
            metadata['sheet_names'] = wb.sheetnames
            
            wb.close()
            
        except Exception as e:
            logger.warning(f"Excel metadata extraction error: {str(e)}")
        
        return metadata
    
    def _extract_ole_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from legacy Office files using OLE."""
        metadata = {}
        try:
            if olefile.isOleFile(file_path):
                ole = olefile.OleFileIO(file_path)
                meta = ole.get_metadata()
                
                metadata['author'] = meta.author
                metadata['last_modified_by'] = meta.last_saved_by
                metadata['creation_date'] = meta.create_time.isoformat() if meta.create_time else None
                metadata['modification_date'] = meta.last_saved_time.isoformat() if meta.last_saved_time else None
                metadata['title'] = meta.title
                metadata['subject'] = meta.subject
                metadata['company'] = meta.company
                metadata['application'] = meta.creating_application
                metadata['revision_count'] = meta.revision_number
                
                ole.close()
                
        except Exception as e:
            logger.warning(f"OLE metadata extraction error: {str(e)}")
        
        return metadata
    
    def _detect_pii_in_metadata(self, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Detect PII in document metadata."""
        pii_findings = []
        
        pii_fields = ['author', 'creator', 'last_modified_by', 'company']
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        
        for field in pii_fields:
            value = metadata.get(field)
            if value and isinstance(value, str) and value.strip():
                if re.search(email_pattern, value):
                    pii_findings.append({
                        'type': 'EMAIL_IN_METADATA',
                        'field': field,
                        'value': value,
                        'risk_level': 'Medium',
                        'description': f'Email address found in document metadata ({field})'
                    })
                elif len(value.split()) >= 2 and value[0].isupper():
                    pii_findings.append({
                        'type': 'NAME_IN_METADATA',
                        'field': field,
                        'value': value,
                        'risk_level': 'Low',
                        'description': f'Potential personal name found in document metadata ({field})'
                    })
        
        return pii_findings
    
    def _detect_tampering(self, file_path: str, ext: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Detect signs of document tampering or integrity issues.
        """
        integrity = {
            'tampering_detected': False,
            'integrity_score': 100,
            'issues': [],
            'file_hash': None,
            'timestamp_consistency': True,
            'metadata_consistency': True
        }
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                integrity['file_hash'] = hashlib.sha256(content).hexdigest()
            
            file_mtime = os.path.getmtime(file_path)
            file_ctime = os.path.getctime(file_path)
            
            if metadata.get('modification_date'):
                try:
                    meta_mtime = datetime.fromisoformat(metadata['modification_date'].replace('Z', '+00:00'))
                    file_mtime_dt = datetime.fromtimestamp(file_mtime)
                    
                    time_diff = abs((meta_mtime.replace(tzinfo=None) - file_mtime_dt).total_seconds())
                    if time_diff > 86400:
                        integrity['issues'].append({
                            'type': 'TIMESTAMP_MISMATCH',
                            'description': 'Document metadata modification date differs significantly from file system timestamp',
                            'severity': 'Medium',
                            'details': f'Metadata: {metadata["modification_date"]}, File system: {file_mtime_dt.isoformat()}'
                        })
                        integrity['timestamp_consistency'] = False
                        integrity['integrity_score'] -= 15
                except:
                    pass
            
            if metadata.get('creation_date') and metadata.get('modification_date'):
                try:
                    creation = datetime.fromisoformat(metadata['creation_date'].replace('Z', '+00:00'))
                    modification = datetime.fromisoformat(metadata['modification_date'].replace('Z', '+00:00'))
                    
                    if modification < creation:
                        integrity['issues'].append({
                            'type': 'TIMESTAMP_ANOMALY',
                            'description': 'Document modification date is earlier than creation date',
                            'severity': 'High',
                            'details': f'Creation: {metadata["creation_date"]}, Modification: {metadata["modification_date"]}'
                        })
                        integrity['metadata_consistency'] = False
                        integrity['integrity_score'] -= 25
                except:
                    pass
            
            if ext == '.pdf' and PDF_AVAILABLE:
                integrity.update(self._check_pdf_integrity(file_path))
            
            if metadata.get('producer') and metadata.get('creator'):
                producer = str(metadata.get('producer', '')).lower()
                creator = str(metadata.get('creator', '')).lower()
                
                suspicious_patterns = ['pdf editor', 'foxit', 'nitro', 'pdf-xchange']
                for pattern in suspicious_patterns:
                    if pattern in producer or pattern in creator:
                        integrity['issues'].append({
                            'type': 'EDITING_SOFTWARE_DETECTED',
                            'description': f'Document was processed with editing software: {pattern}',
                            'severity': 'Low',
                            'details': f'Producer: {metadata.get("producer")}, Creator: {metadata.get("creator")}'
                        })
                        integrity['integrity_score'] -= 5
                        break
            
            integrity['tampering_detected'] = len(integrity['issues']) > 0
            integrity['integrity_score'] = max(0, integrity['integrity_score'])
            
        except Exception as e:
            logger.warning(f"Tampering detection error: {str(e)}")
            integrity['error'] = str(e)
        
        return integrity
    
    def _check_pdf_integrity(self, file_path: str) -> Dict[str, Any]:
        """Check PDF-specific integrity indicators."""
        result = {'incremental_updates': 0}
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                
                eof_count = content.count(b'%%EOF')
                if eof_count > 1:
                    result['incremental_updates'] = eof_count - 1
                    result['issues'] = result.get('issues', [])
                    result['issues'].append({
                        'type': 'INCREMENTAL_UPDATES',
                        'description': f'PDF has {eof_count - 1} incremental update(s), indicating modifications',
                        'severity': 'Low',
                        'details': f'Number of %%EOF markers: {eof_count}'
                    })
        except:
            pass
        return result
    
    def _detect_language(self, text: str) -> Dict[str, Any]:
        """
        Detect document language for proper PII pattern matching.
        """
        result = {
            'detected_language': None,
            'language_code': None,
            'confidence': 0.0,
            'all_languages': [],
            'is_multilingual': False
        }
        
        if not text or len(text.strip()) < 50:
            result['detected_language'] = 'unknown'
            return result
        
        if not LANGDETECT_AVAILABLE:
            result['detected_language'] = self._detect_language_heuristic(text)
            result['confidence'] = 0.7
            return result
        
        try:
            sample_text = text[:5000]
            
            langs = detect_langs(sample_text)
            if langs:
                primary = langs[0]
                result['detected_language'] = self._get_language_name(primary.lang)
                result['language_code'] = primary.lang
                result['confidence'] = round(primary.prob, 3)
                
                result['all_languages'] = [
                    {'language': self._get_language_name(l.lang), 'code': l.lang, 'probability': round(l.prob, 3)}
                    for l in langs[:5]
                ]
                
                result['is_multilingual'] = len([l for l in langs if l.prob > 0.2]) > 1
                
        except Exception as e:
            logger.warning(f"Language detection error: {str(e)}")
            result['detected_language'] = 'unknown'
            result['error'] = str(e)
        
        return result
    
    def _detect_language_heuristic(self, text: str) -> str:
        """Simple heuristic language detection."""
        dutch_words = ['de', 'het', 'een', 'van', 'en', 'dat', 'is', 'op', 'te', 'voor', 'met', 'zijn', 'niet', 'aan', 'wordt']
        german_words = ['der', 'die', 'das', 'und', 'ist', 'von', 'mit', 'auf', 'für', 'nicht', 'sich', 'auch', 'werden']
        french_words = ['le', 'la', 'les', 'de', 'et', 'est', 'un', 'une', 'que', 'pour', 'dans', 'pas', 'sur', 'avec']
        
        text_lower = text.lower()
        words = text_lower.split()
        
        dutch_count = sum(1 for w in words if w in dutch_words)
        german_count = sum(1 for w in words if w in german_words)
        french_count = sum(1 for w in words if w in french_words)
        english_count = len(words) - dutch_count - german_count - french_count
        
        counts = {'Dutch': dutch_count, 'German': german_count, 'French': french_count, 'English': english_count}
        return max(counts, key=counts.get)
    
    def _get_language_name(self, code: str) -> str:
        """Convert language code to name."""
        language_map = {
            'nl': 'Dutch', 'en': 'English', 'de': 'German', 'fr': 'French',
            'es': 'Spanish', 'it': 'Italian', 'pt': 'Portuguese', 'pl': 'Polish',
            'ru': 'Russian', 'zh-cn': 'Chinese', 'ja': 'Japanese', 'ko': 'Korean'
        }
        return language_map.get(code, code.upper())
    
    def _classify_document(self, text: str, file_name: str, ext: str) -> Dict[str, Any]:
        """
        Automatically classify document type.
        """
        result = {
            'primary_category': 'general',
            'confidence': 0.0,
            'all_categories': [],
            'sensitivity_level': 'low',
            'gdpr_relevance': 'low'
        }
        
        text_lower = text.lower()
        file_lower = file_name.lower()
        
        scores = {}
        for category, keywords in self.classification_keywords.items():
            score = 0
            matches = []
            for keyword in keywords:
                if keyword in text_lower:
                    score += text_lower.count(keyword)
                    matches.append(keyword)
                if keyword in file_lower:
                    score += 5
            scores[category] = {'score': score, 'matches': matches}
        
        if scores:
            total_score = sum(s['score'] for s in scores.values())
            if total_score > 0:
                sorted_scores = sorted(scores.items(), key=lambda x: x[1]['score'], reverse=True)
                
                result['primary_category'] = sorted_scores[0][0]
                result['confidence'] = round(sorted_scores[0][1]['score'] / total_score, 3) if total_score > 0 else 0
                result['matched_keywords'] = sorted_scores[0][1]['matches'][:10]
                
                result['all_categories'] = [
                    {'category': cat, 'score': data['score'], 'confidence': round(data['score'] / total_score, 3)}
                    for cat, data in sorted_scores[:5] if data['score'] > 0
                ]
        
        high_sensitivity = ['medical', 'hr', 'legal', 'financial', 'personal']
        medium_sensitivity = ['contract', 'invoice']
        
        if result['primary_category'] in high_sensitivity:
            result['sensitivity_level'] = 'high'
            result['gdpr_relevance'] = 'high'
        elif result['primary_category'] in medium_sensitivity:
            result['sensitivity_level'] = 'medium'
            result['gdpr_relevance'] = 'medium'
        
        return result
    
    def _detect_hidden_content(self, file_path: str, ext: str) -> Dict[str, Any]:
        """
        Detect hidden content in documents.
        """
        result = {
            'hidden_content_found': False,
            'hidden_items': [],
            'hidden_text': False,
            'hidden_layers': False,
            'hidden_rows_columns': False,
            'hidden_sheets': False,
            'comments': []
        }
        
        try:
            if ext == '.pdf' and PDF_AVAILABLE:
                result.update(self._detect_pdf_hidden_content(file_path))
            elif ext in ['.xlsx', '.xls'] and EXCEL_AVAILABLE:
                result.update(self._detect_excel_hidden_content(file_path))
            elif ext == '.docx' and DOCX_AVAILABLE:
                result.update(self._detect_docx_hidden_content(file_path))
                
        except Exception as e:
            logger.warning(f"Hidden content detection error: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _detect_pdf_hidden_content(self, file_path: str) -> Dict[str, Any]:
        """Detect hidden content in PDFs."""
        result = {'hidden_items': []}
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                for i, page in enumerate(reader.pages):
                    if '/Annots' in page:
                        annots = page['/Annots']
                        if annots:
                            result['hidden_items'].append({
                                'type': 'ANNOTATIONS',
                                'page': i + 1,
                                'description': 'Page contains annotations that may hide content'
                            })
                    
                    if '/OCProperties' in reader.trailer.get('/Root', {}):
                        result['hidden_layers'] = True
                        result['hidden_items'].append({
                            'type': 'OPTIONAL_CONTENT_LAYERS',
                            'description': 'PDF contains optional content layers (may contain hidden content)'
                        })
                
                if reader.trailer.get('/Root', {}).get('/AcroForm'):
                    result['hidden_items'].append({
                        'type': 'FORM_FIELDS',
                        'description': 'PDF contains form fields that may contain hidden data'
                    })
                    
        except Exception as e:
            logger.warning(f"PDF hidden content detection error: {str(e)}")
        
        result['hidden_content_found'] = len(result['hidden_items']) > 0
        return result
    
    def _detect_excel_hidden_content(self, file_path: str) -> Dict[str, Any]:
        """Detect hidden content in Excel files."""
        result = {'hidden_items': [], 'comments': []}
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet in wb.worksheets:
                if sheet.sheet_state == 'hidden':
                    result['hidden_sheets'] = True
                    result['hidden_items'].append({
                        'type': 'HIDDEN_SHEET',
                        'sheet': sheet.title,
                        'description': f'Hidden worksheet: {sheet.title}'
                    })
                
                hidden_rows = [r for r in sheet.row_dimensions if sheet.row_dimensions[r].hidden]
                if hidden_rows:
                    result['hidden_rows_columns'] = True
                    result['hidden_items'].append({
                        'type': 'HIDDEN_ROWS',
                        'sheet': sheet.title,
                        'count': len(hidden_rows),
                        'description': f'{len(hidden_rows)} hidden rows in sheet {sheet.title}'
                    })
                
                hidden_cols = [c for c in sheet.column_dimensions if sheet.column_dimensions[c].hidden]
                if hidden_cols:
                    result['hidden_rows_columns'] = True
                    result['hidden_items'].append({
                        'type': 'HIDDEN_COLUMNS',
                        'sheet': sheet.title,
                        'count': len(hidden_cols),
                        'description': f'{len(hidden_cols)} hidden columns in sheet {sheet.title}'
                    })
                
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.comment:
                            result['comments'].append({
                                'sheet': sheet.title,
                                'cell': cell.coordinate,
                                'author': cell.comment.author if hasattr(cell.comment, 'author') else None,
                                'text': cell.comment.text[:100] if cell.comment.text else None
                            })
            
            wb.close()
            
        except Exception as e:
            logger.warning(f"Excel hidden content detection error: {str(e)}")
        
        result['hidden_content_found'] = len(result['hidden_items']) > 0 or len(result['comments']) > 0
        return result
    
    def _detect_docx_hidden_content(self, file_path: str) -> Dict[str, Any]:
        """Detect hidden content in DOCX files."""
        result = {'hidden_items': [], 'comments': []}
        try:
            doc = docx.Document(file_path)
            
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.hidden:
                        result['hidden_text'] = True
                        result['hidden_items'].append({
                            'type': 'HIDDEN_TEXT',
                            'description': 'Document contains hidden text',
                            'sample': run.text[:50] if run.text else None
                        })
                        break
            
            try:
                if hasattr(doc, 'comments'):
                    for comment in doc.comments:
                        result['comments'].append({
                            'author': comment.author,
                            'text': comment.text[:100] if comment.text else None
                        })
            except:
                pass
                
        except Exception as e:
            logger.warning(f"DOCX hidden content detection error: {str(e)}")
        
        result['hidden_content_found'] = len(result['hidden_items']) > 0
        return result
    
    def _detect_embedded_objects(self, file_path: str, ext: str) -> Dict[str, Any]:
        """
        Detect embedded objects, macros, and scripts.
        """
        result = {
            'embedded_objects_found': False,
            'objects': [],
            'macros_detected': False,
            'scripts_detected': False,
            'ole_objects': [],
            'risk_level': 'low'
        }
        
        try:
            if ext in ['.doc', '.xls', '.ppt'] and OLE_AVAILABLE:
                result.update(self._detect_ole_embedded(file_path))
            elif ext == '.docx' and DOCX_AVAILABLE:
                result.update(self._detect_docx_embedded(file_path))
            elif ext == '.xlsx' and EXCEL_AVAILABLE:
                result.update(self._detect_excel_embedded(file_path))
            elif ext == '.pdf' and PDF_AVAILABLE:
                result.update(self._detect_pdf_embedded(file_path))
                
        except Exception as e:
            logger.warning(f"Embedded object detection error: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _detect_ole_embedded(self, file_path: str) -> Dict[str, Any]:
        """Detect embedded objects in OLE files."""
        result = {'objects': [], 'ole_objects': []}
        try:
            if olefile.isOleFile(file_path):
                ole = olefile.OleFileIO(file_path)
                
                vba_streams = ['Macros', 'VBA', '_VBA_PROJECT_CUR', 'VBA_PROJECT']
                for stream_name in vba_streams:
                    if ole.exists(stream_name):
                        result['macros_detected'] = True
                        result['objects'].append({
                            'type': 'VBA_MACRO',
                            'name': stream_name,
                            'description': 'VBA macro code detected',
                            'risk': 'high'
                        })
                
                for stream in ole.listdir():
                    stream_path = '/'.join(stream)
                    if 'ObjectPool' in stream_path or 'Embedding' in stream_path:
                        result['ole_objects'].append({
                            'type': 'EMBEDDED_OBJECT',
                            'path': stream_path,
                            'description': 'Embedded OLE object detected'
                        })
                
                ole.close()
                
        except Exception as e:
            logger.warning(f"OLE embedded detection error: {str(e)}")
        
        result['embedded_objects_found'] = len(result['objects']) > 0 or len(result['ole_objects']) > 0
        if result['macros_detected']:
            result['risk_level'] = 'high'
        elif result['embedded_objects_found']:
            result['risk_level'] = 'medium'
        
        return result
    
    def _detect_docx_embedded(self, file_path: str) -> Dict[str, Any]:
        """Detect embedded objects in DOCX files."""
        result = {'objects': []}
        try:
            doc = docx.Document(file_path)
            
            for rel in doc.part.rels.values():
                if 'oleObject' in str(rel.reltype).lower():
                    result['objects'].append({
                        'type': 'OLE_OBJECT',
                        'relationship_type': str(rel.reltype),
                        'description': 'Embedded OLE object in document'
                    })
                elif 'image' in str(rel.reltype).lower():
                    result['objects'].append({
                        'type': 'EMBEDDED_IMAGE',
                        'relationship_type': str(rel.reltype),
                        'description': 'Embedded image in document'
                    })
                    
        except Exception as e:
            logger.warning(f"DOCX embedded detection error: {str(e)}")
        
        result['embedded_objects_found'] = len(result['objects']) > 0
        return result
    
    def _detect_excel_embedded(self, file_path: str) -> Dict[str, Any]:
        """Detect embedded objects in Excel files."""
        result = {'objects': []}
        try:
            wb = openpyxl.load_workbook(file_path)
            
            if hasattr(wb, 'vba_archive') and wb.vba_archive:
                result['macros_detected'] = True
                result['objects'].append({
                    'type': 'VBA_MACRO',
                    'description': 'VBA macros detected in workbook',
                    'risk': 'high'
                })
                result['risk_level'] = 'high'
            
            wb.close()
            
        except Exception as e:
            logger.warning(f"Excel embedded detection error: {str(e)}")
        
        result['embedded_objects_found'] = len(result['objects']) > 0
        return result
    
    def _detect_pdf_embedded(self, file_path: str) -> Dict[str, Any]:
        """Detect embedded files in PDFs."""
        result = {'objects': []}
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                if '/Names' in reader.trailer.get('/Root', {}):
                    root = reader.trailer['/Root']
                    if '/EmbeddedFiles' in root.get('/Names', {}):
                        result['objects'].append({
                            'type': 'EMBEDDED_FILES',
                            'description': 'PDF contains embedded files'
                        })
                
                if '/AA' in reader.trailer.get('/Root', {}) or '/OpenAction' in reader.trailer.get('/Root', {}):
                    result['scripts_detected'] = True
                    result['objects'].append({
                        'type': 'JAVASCRIPT_ACTION',
                        'description': 'PDF contains JavaScript or automatic actions',
                        'risk': 'high'
                    })
                    result['risk_level'] = 'high'
                    
        except Exception as e:
            logger.warning(f"PDF embedded detection error: {str(e)}")
        
        result['embedded_objects_found'] = len(result['objects']) > 0
        return result
    
    def _detect_digital_signatures(self, file_path: str, ext: str) -> Dict[str, Any]:
        """
        Detect and analyze digital signatures.
        """
        result = {
            'has_signature': False,
            'signatures': [],
            'signature_valid': None,
            'signer_info': None
        }
        
        try:
            if ext == '.pdf' and PDF_AVAILABLE:
                result.update(self._detect_pdf_signatures(file_path))
            elif ext == '.docx' and DOCX_AVAILABLE:
                result.update(self._detect_docx_signatures(file_path))
                
        except Exception as e:
            logger.warning(f"Digital signature detection error: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _detect_pdf_signatures(self, file_path: str) -> Dict[str, Any]:
        """Detect digital signatures in PDFs."""
        result = {'signatures': []}
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                if '/AcroForm' in reader.trailer.get('/Root', {}):
                    acroform = reader.trailer['/Root']['/AcroForm']
                    if '/SigFlags' in acroform:
                        result['has_signature'] = True
                        result['signatures'].append({
                            'type': 'ACROFORM_SIGNATURE',
                            'description': 'PDF contains AcroForm digital signature field'
                        })
                        
        except Exception as e:
            logger.warning(f"PDF signature detection error: {str(e)}")
        
        return result
    
    def _detect_docx_signatures(self, file_path: str) -> Dict[str, Any]:
        """Detect digital signatures in DOCX files."""
        result = {'signatures': []}
        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as z:
                if '_xmlsignatures/sig1.xml' in z.namelist() or 'docProps/sig1.xml' in z.namelist():
                    result['has_signature'] = True
                    result['signatures'].append({
                        'type': 'XML_SIGNATURE',
                        'description': 'Document contains XML digital signature'
                    })
                    
        except Exception as e:
            logger.warning(f"DOCX signature detection error: {str(e)}")
        
        return result
    
    def _detect_redaction(self, text: str, file_path: str, ext: str) -> Dict[str, Any]:
        """
        Detect redacted content in documents.
        """
        result = {
            'redaction_detected': False,
            'redaction_patterns': [],
            'potential_hidden_pii': False
        }
        
        try:
            redaction_patterns = [
                (r'█+', 'BLACK_BLOCK_CHARS'),
                (r'▓+', 'DARK_SHADE_CHARS'),
                (r'▒+', 'MEDIUM_SHADE_CHARS'),
                (r'\[REDACTED\]', 'REDACTED_LABEL'),
                (r'\[CONFIDENTIAL\]', 'CONFIDENTIAL_LABEL'),
                (r'\[REMOVED\]', 'REMOVED_LABEL'),
                (r'X{5,}', 'X_PATTERN'),
                (r'\*{5,}', 'ASTERISK_PATTERN'),
                (r'_{5,}', 'UNDERSCORE_PATTERN')
            ]
            
            for pattern, pattern_type in redaction_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    result['redaction_detected'] = True
                    result['redaction_patterns'].append({
                        'type': pattern_type,
                        'count': len(matches),
                        'description': f'Detected {len(matches)} instance(s) of {pattern_type} redaction pattern'
                    })
            
            if result['redaction_detected']:
                result['potential_hidden_pii'] = True
                
        except Exception as e:
            logger.warning(f"Redaction detection error: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _analyze_version_history(self, file_path: str, ext: str) -> Dict[str, Any]:
        """
        Analyze document version/revision history.
        """
        result = {
            'has_revisions': False,
            'revision_count': 0,
            'revisions': [],
            'track_changes': False
        }
        
        try:
            if ext == '.docx' and DOCX_AVAILABLE:
                result.update(self._analyze_docx_revisions(file_path))
            elif ext == '.pdf' and PDF_AVAILABLE:
                result.update(self._analyze_pdf_revisions(file_path))
                
        except Exception as e:
            logger.warning(f"Version history analysis error: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _analyze_docx_revisions(self, file_path: str) -> Dict[str, Any]:
        """Analyze DOCX revision history."""
        result = {'revisions': []}
        try:
            doc = docx.Document(file_path)
            
            if doc.core_properties.revision:
                result['revision_count'] = int(doc.core_properties.revision) if doc.core_properties.revision else 0
                result['has_revisions'] = result['revision_count'] > 1
            
            result['revisions'].append({
                'type': 'DOCUMENT_INFO',
                'revision_number': doc.core_properties.revision,
                'last_modified_by': doc.core_properties.last_modified_by,
                'modified_date': doc.core_properties.modified.isoformat() if doc.core_properties.modified else None
            })
                
        except Exception as e:
            logger.warning(f"DOCX revision analysis error: {str(e)}")
        
        return result
    
    def _analyze_pdf_revisions(self, file_path: str) -> Dict[str, Any]:
        """Analyze PDF revision history (incremental saves)."""
        result = {'revisions': []}
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                
                eof_positions = []
                pos = 0
                while True:
                    pos = content.find(b'%%EOF', pos)
                    if pos == -1:
                        break
                    eof_positions.append(pos)
                    pos += 1
                
                if len(eof_positions) > 1:
                    result['has_revisions'] = True
                    result['revision_count'] = len(eof_positions)
                    result['revisions'].append({
                        'type': 'INCREMENTAL_UPDATES',
                        'count': len(eof_positions),
                        'description': f'PDF has {len(eof_positions)} revision(s) through incremental updates'
                    })
                    
        except Exception as e:
            logger.warning(f"PDF revision analysis error: {str(e)}")
        
        return result
    
    def _extract_table_data(self, file_path: str, ext: str) -> Dict[str, Any]:
        """
        Extract structured table data from documents.
        """
        result = {
            'tables_found': 0,
            'tables': [],
            'forms_found': 0,
            'form_fields': []
        }
        
        try:
            if ext == '.docx' and DOCX_AVAILABLE:
                result.update(self._extract_docx_tables(file_path))
            elif ext in ['.xlsx', '.xls'] and EXCEL_AVAILABLE:
                result.update(self._extract_excel_tables(file_path))
            elif ext == '.pdf' and PDF_AVAILABLE:
                result.update(self._extract_pdf_forms(file_path))
                
        except Exception as e:
            logger.warning(f"Table extraction error: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _extract_docx_tables(self, file_path: str) -> Dict[str, Any]:
        """Extract tables from DOCX files."""
        result = {'tables': []}
        try:
            doc = docx.Document(file_path)
            
            for i, table in enumerate(doc.tables):
                table_info = {
                    'index': i + 1,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'headers': []
                }
                
                if table.rows:
                    first_row = table.rows[0]
                    table_info['headers'] = [cell.text[:50] for cell in first_row.cells]
                
                result['tables'].append(table_info)
            
            result['tables_found'] = len(doc.tables)
            
        except Exception as e:
            logger.warning(f"DOCX table extraction error: {str(e)}")
        
        return result
    
    def _extract_excel_tables(self, file_path: str) -> Dict[str, Any]:
        """Extract table information from Excel files."""
        result = {'tables': []}
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True)
            
            for sheet in wb.worksheets:
                table_info = {
                    'sheet': sheet.title,
                    'max_row': sheet.max_row,
                    'max_column': sheet.max_column,
                    'headers': []
                }
                
                if sheet.max_row and sheet.max_row > 0:
                    first_row = list(sheet.iter_rows(min_row=1, max_row=1))[0] if sheet.max_row > 0 else []
                    table_info['headers'] = [str(cell.value)[:50] if cell.value else '' for cell in first_row[:20]]
                
                result['tables'].append(table_info)
            
            result['tables_found'] = len(wb.worksheets)
            wb.close()
            
        except Exception as e:
            logger.warning(f"Excel table extraction error: {str(e)}")
        
        return result
    
    def _extract_pdf_forms(self, file_path: str) -> Dict[str, Any]:
        """Extract form fields from PDFs."""
        result = {'form_fields': []}
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                if '/AcroForm' in reader.trailer.get('/Root', {}):
                    acroform = reader.trailer['/Root']['/AcroForm']
                    if '/Fields' in acroform:
                        fields = acroform['/Fields']
                        result['forms_found'] = len(fields) if isinstance(fields, list) else 1
                        result['form_fields'].append({
                            'type': 'ACROFORM',
                            'field_count': result['forms_found'],
                            'description': 'PDF contains interactive form fields'
                        })
                        
        except Exception as e:
            logger.warning(f"PDF form extraction error: {str(e)}")
        
        return result
    
    def _perform_forensic_analysis(self, file_path: str, ext: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Perform document forensic analysis.
        """
        result = {
            'creation_timeline': {},
            'authorship_analysis': {},
            'anomalies': [],
            'forensic_score': 100
        }
        
        try:
            result['creation_timeline'] = {
                'file_created': metadata.get('file_created'),
                'file_modified': metadata.get('file_modified'),
                'metadata_created': metadata.get('creation_date'),
                'metadata_modified': metadata.get('modification_date'),
                'timeline_consistent': True
            }
            
            dates = [
                ('file_created', metadata.get('file_created')),
                ('file_modified', metadata.get('file_modified')),
                ('metadata_created', metadata.get('creation_date')),
                ('metadata_modified', metadata.get('modification_date'))
            ]
            
            valid_dates = [(name, d) for name, d in dates if d]
            if len(valid_dates) >= 2:
                parsed_dates = []
                for name, d in valid_dates:
                    try:
                        parsed = datetime.fromisoformat(d.replace('Z', '+00:00'))
                        parsed_dates.append((name, parsed))
                    except:
                        pass
                
                if parsed_dates:
                    sorted_dates = sorted(parsed_dates, key=lambda x: x[1])
                    result['creation_timeline']['chronological_order'] = [name for name, _ in sorted_dates]
            
            result['authorship_analysis'] = {
                'author': metadata.get('author'),
                'creator': metadata.get('creator'),
                'last_modified_by': metadata.get('last_modified_by'),
                'company': metadata.get('company'),
                'application': metadata.get('producer') or metadata.get('application'),
                'multiple_authors': metadata.get('author') != metadata.get('last_modified_by') if metadata.get('author') and metadata.get('last_modified_by') else False
            }
            
            if metadata.get('author') and metadata.get('last_modified_by'):
                if metadata['author'] != metadata['last_modified_by']:
                    result['anomalies'].append({
                        'type': 'MULTIPLE_AUTHORS',
                        'description': 'Document was modified by a different user than the original author',
                        'severity': 'Info',
                        'author': metadata['author'],
                        'last_modified_by': metadata['last_modified_by']
                    })
            
            if metadata.get('file_size', 0) < 100:
                result['anomalies'].append({
                    'type': 'SUSPICIOUSLY_SMALL',
                    'description': 'Document file size is unusually small',
                    'severity': 'Low'
                })
                result['forensic_score'] -= 10
            
        except Exception as e:
            logger.warning(f"Forensic analysis error: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _detect_cross_references(self, text: str, file_path: str) -> Dict[str, Any]:
        """
        Detect cross-references to other documents or external resources.
        """
        result = {
            'references_found': 0,
            'internal_references': [],
            'external_references': [],
            'urls': [],
            'file_references': [],
            'data_sharing_indicators': []
        }
        
        try:
            url_pattern = r'https?://[^\s<>"\']+|www\.[^\s<>"\']+'
            urls = re.findall(url_pattern, text, re.IGNORECASE)
            result['urls'] = list(set(urls))[:50]
            
            file_patterns = [
                r'[A-Za-z]:\\[^\s<>"\']+\.[a-z]{2,4}',
                r'/(?:home|Users|var|opt)/[^\s<>"\']+\.[a-z]{2,4}',
                r'\\\\[^\s<>"\']+\.[a-z]{2,4}',
                r'["\'][^"\']+\.(pdf|docx?|xlsx?|csv|txt|json|xml)["\']'
            ]
            
            for pattern in file_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                result['file_references'].extend(matches[:20])
            
            result['file_references'] = list(set(result['file_references']))[:30]
            
            sharing_patterns = [
                (r'shared\s+with|sent\s+to|forwarded\s+to|cc[:\s]+|bcc[:\s]+', 'EMAIL_SHARING'),
                (r'uploaded\s+to|synced\s+with|backed\s+up\s+to', 'CLOUD_SHARING'),
                (r'transferred\s+to|exported\s+to|copied\s+to', 'DATA_TRANSFER'),
                (r'see\s+also|refer\s+to|as\s+per|according\s+to', 'DOCUMENT_REFERENCE')
            ]
            
            for pattern, indicator_type in sharing_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    result['data_sharing_indicators'].append({
                        'type': indicator_type,
                        'count': len(matches),
                        'description': f'Found {len(matches)} {indicator_type.replace("_", " ").lower()} indicator(s)'
                    })
            
            result['references_found'] = len(result['urls']) + len(result['file_references'])
            
        except Exception as e:
            logger.warning(f"Cross-reference detection error: {str(e)}")
            result['error'] = str(e)
        
        return result
    
    def _generate_findings(self, analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Generate findings from analysis results.
        """
        findings = []
        
        if analysis.get('metadata', {}).get('pii_in_metadata'):
            for pii in analysis['metadata']['pii_in_metadata']:
                findings.append({
                    'type': pii['type'],
                    'category': 'METADATA_PII',
                    'description': pii['description'],
                    'value': pii.get('value', ''),
                    'risk_level': pii.get('risk_level', 'Medium'),
                    'source': 'document_metadata',
                    'recommendation': 'Review and remove personal information from document metadata before sharing'
                })
        
        if analysis.get('integrity', {}).get('issues'):
            for issue in analysis['integrity']['issues']:
                findings.append({
                    'type': issue['type'],
                    'category': 'INTEGRITY',
                    'description': issue['description'],
                    'details': issue.get('details', ''),
                    'risk_level': issue.get('severity', 'Medium'),
                    'source': 'integrity_check',
                    'recommendation': 'Verify document authenticity and review modification history'
                })
        
        if analysis.get('hidden_content', {}).get('hidden_content_found'):
            for item in analysis['hidden_content'].get('hidden_items', []):
                findings.append({
                    'type': item['type'],
                    'category': 'HIDDEN_CONTENT',
                    'description': item['description'],
                    'risk_level': 'Medium',
                    'source': 'hidden_content_detection',
                    'recommendation': 'Review hidden content for sensitive information before sharing'
                })
        
        if analysis.get('embedded_objects', {}).get('macros_detected'):
            findings.append({
                'type': 'MACRO_DETECTED',
                'category': 'SECURITY',
                'description': 'Document contains VBA macros which may pose security risks',
                'risk_level': 'High',
                'source': 'embedded_object_detection',
                'recommendation': 'Review macro code for malicious content and consider disabling macros'
            })
        
        if analysis.get('redaction', {}).get('redaction_detected'):
            findings.append({
                'type': 'REDACTION_DETECTED',
                'category': 'PRIVACY',
                'description': 'Document contains redacted content that may hide sensitive information',
                'risk_level': 'Medium',
                'source': 'redaction_detection',
                'recommendation': 'Verify redaction is complete and cannot be reversed'
            })
        
        if analysis.get('classification', {}).get('sensitivity_level') == 'high':
            findings.append({
                'type': 'HIGH_SENSITIVITY_DOCUMENT',
                'category': 'CLASSIFICATION',
                'description': f"Document classified as {analysis['classification'].get('primary_category', 'unknown')} with high sensitivity",
                'risk_level': 'High',
                'source': 'document_classification',
                'recommendation': 'Apply appropriate access controls and handling procedures'
            })
        
        if analysis.get('cross_references', {}).get('urls'):
            external_urls = [u for u in analysis['cross_references']['urls'] if not any(safe in u.lower() for safe in ['example.com', 'localhost'])]
            if len(external_urls) > 5:
                findings.append({
                    'type': 'EXTERNAL_REFERENCES',
                    'category': 'DATA_SHARING',
                    'description': f'Document contains {len(external_urls)} external URL references',
                    'risk_level': 'Low',
                    'source': 'cross_reference_detection',
                    'recommendation': 'Review external links for data sharing implications'
                })
        
        return findings
    
    def _calculate_risk_indicators(self, analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Calculate overall risk indicators.
        """
        indicators = []
        
        if analysis.get('integrity', {}).get('integrity_score', 100) < 70:
            indicators.append({
                'indicator': 'LOW_INTEGRITY_SCORE',
                'value': analysis['integrity']['integrity_score'],
                'severity': 'High',
                'description': 'Document integrity concerns detected'
            })
        
        if analysis.get('embedded_objects', {}).get('risk_level') == 'high':
            indicators.append({
                'indicator': 'HIGH_RISK_EMBEDDED_CONTENT',
                'severity': 'High',
                'description': 'Document contains potentially dangerous embedded content'
            })
        
        if analysis.get('classification', {}).get('gdpr_relevance') == 'high':
            indicators.append({
                'indicator': 'HIGH_GDPR_RELEVANCE',
                'severity': 'Medium',
                'description': 'Document likely contains GDPR-relevant personal data'
            })
        
        if analysis.get('hidden_content', {}).get('hidden_content_found'):
            indicators.append({
                'indicator': 'HIDDEN_CONTENT_PRESENT',
                'severity': 'Medium',
                'description': 'Document contains hidden content that may include sensitive data'
            })
        
        return indicators


def create_document_analyzer(region: str = "Netherlands") -> AdvancedDocumentAnalyzer:
    """Factory function to create AdvancedDocumentAnalyzer instance."""
    return AdvancedDocumentAnalyzer(region=region)
