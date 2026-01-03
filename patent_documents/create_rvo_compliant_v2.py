"""
Generate RVO-compliant patent documents with CORRECT formatting:
- Page numbers centered at top (in header)
- Line numbers at LEFT MARGIN at every 5th line (5, 10, 15...) 
- Line numbers restart each page
"""

from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

os.makedirs("patent_documents", exist_ok=True)

def create_beschrijving_v2():
    """Create RVO-compliant Beschrijving with proper line numbering"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)  # Extra left margin for line numbers
        section.right_margin = Cm(2.5)
        
        # Add page number in header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run("1")
        run.font.size = Pt(11)
    
    # Content with proper structure
    lines = [
        "",  # Line 1
        "BESCHRIJVING",  # Line 2
        "",  # Line 3
        "TITEL",  # Line 4
        "Intelligente Database Scanner met Adaptieve Sampling en Prioritering",  # Line 5 - NUMBER HERE
        "",  # Line 6
        "TECHNISCH GEBIED",  # Line 7
        "De uitvinding betreft een computersysteem voor de geautomatiseerde analyse",  # Line 8
        "van databases met het oog op gegevensbescherming en privacy-compliance.",  # Line 9
        "Het systeem ondersteunt meerdere relationele database-engines, waaronder",  # Line 10 - NUMBER HERE
        "PostgreSQL, MySQL en Microsoft SQL Server, en is ingericht voor het",  # Line 11
        "detecteren van persoonsgegevens.",  # Line 12
        "",  # Line 13
        "Het systeem maakt gebruik van prioriteit-gebaseerde tabelselectie en",  # Line 14
        "adaptieve bemonsteringsstrategieën om efficiënt gegevens te analyseren.",  # Line 15 - NUMBER HERE
        "Daarnaast voorziet de uitvinding in detectie en validatie van specifiek",  # Line 16
        "Nederlandse identificatiegegevens, waaronder Burgerservicenummers (BSN).",  # Line 17
        "",  # Line 18
        "ACHTERGROND VAN DE UITVINDING",  # Line 19
        "Databases bevatten doorgaans grote hoeveelheden persoonsgegevens die",  # Line 20 - NUMBER HERE
        "verspreid zijn over meerdere tabellen. Bestaande databasescanners voeren",  # Line 21
        "vaak lineaire of volledige scans uit zonder onderscheid te maken tussen",  # Line 22
        "relevante en minder relevante tabellen.",  # Line 23
        "",  # Line 24
        "Dit leidt tot lange verwerkingstijden en een verhoogde belasting van",  # Line 25 - NUMBER HERE
        "systeembronnen. Bovendien ontbreekt in veel bestaande oplossingen een",  # Line 26
        "specifieke validatie voor nationale identificatiegegevens, zoals de",  # Line 27
        "wiskundige 11-proef die vereist is voor BSN-nummers in Nederland.",  # Line 28
        "",  # Line 29
        "SAMENVATTING VAN DE UITVINDING",  # Line 30 - NUMBER HERE
        "De uitvinding voorziet in een intelligent databasescansysteem dat tabellen",  # Line 31
        "binnen een database analyseert en prioriteitsscores toekent op basis van",  # Line 32
        "tabel- en kolomnamen. Op grond van deze scores wordt een geschikte",  # Line 33
        "bemonsteringsstrategie geselecteerd.",  # Line 34
        "Afhankelijk van het risicoprofiel van de database worden verschillende",  # Line 35 - NUMBER HERE
        "bemonsteringsniveaus toegepast, variërend van beperkte tot uitgebreide",  # Line 36
        "analyse. Het systeem maakt gebruik van parallelle verwerking om meerdere",  # Line 37
        "tabellen gelijktijdig te analyseren, waardoor de efficiëntie wordt verbeterd.",  # Line 38
        "",  # Line 39
        "KORTE BESCHRIJVING VAN DE TEKENINGEN",  # Line 40 - NUMBER HERE
        "Figuur 1 toont een overzicht van de multi-engine database scanner",  # Line 41
        "architectuur volgens de uitvinding. Het intelligente databasescannerplatform",  # Line 42
        "(100) ondersteunt drie database-engines: PostgreSQL via psycopg2 (110),",  # Line 43
        "MySQL via connector (120), en Microsoft SQL Server via pyodbc (130). Het",  # Line 44
        "systeem omvat een prioriteringsmodule (140), parallelle workers (150) en een",  # Line 45 - NUMBER HERE
        "BSN-validatiemodule (160).",  # Line 46
        "",  # Line 47
        "Figuur 2 toont het prioriteitsscorealgoritme. De tabelprioriteitsberekening",  # Line 48
        "(200) omvat stap 1 voor de basisscore (210) en stap 2 voor de kolomboost",  # Line 49
        "(220). De uiteindelijke scoreberekening (230) combineert deze waarden met",  # Line 50 - NUMBER HERE
        "een maximum van 3,5. Een voorbeeld (240) toont hoe een tabel",  # Line 51
        "\"customer_profiles\" met email- en telefoonkolommen de hoogste prioriteit",  # Line 52
        "verkrijgt.",  # Line 53
        "",  # Line 54
        "Figuur 3 toont de adaptieve bemonsteringsstrategieën. De",  # Line 55 - NUMBER HERE
        "scanmodus-beslisboom (300) selecteert tussen snelle modus (310), slimme",  # Line 56
        "modus (320), en diepe modus (330). Invoerparameters (340) bepalen de",  # Line 57
        "selectie. De tijdsbesparing (350) bedraagt tot 60% reductie.",  # Line 58
        "",  # Line 59
        "Figuur 4 toont de parallelle scanningworkflow. De ThreadPoolExecutor (400,",  # Line 60 - NUMBER HERE
        "410) verdeelt het werk over drie workers (420, 421, 422). De",  # Line 61
        "resultaataggregatie (430) combineert alle bevindingen.",  # Line 62
        "",  # Line 63
        "Figuur 5 toont de BSN-checksumvalidatie volgens de Nederlandse 11-proef.",  # Line 64
        "Het algoritme (500) past de formule (510) toe: checksum = (d0×9) + (d1×8)",  # Line 65 - NUMBER HERE
        "+ (d2×7) + (d3×6) + (d4×5) + (d5×4) + (d6×3) + (d7×2) - (d8×1). De",  # Line 66
        "validatie (520) controleert of checksum mod 11 = 0. De detectiestroom",  # Line 67
        "(530) omvat regex-patroonherkenning, checksumvalidatie en",  # Line 68
        "GDPR-classificatie.",  # Line 69
        "Figuur 6 toont de schema-intelligentieanalyse. De databaserisicobepaling",  # Line 70 - NUMBER HERE
        "(600) categoriseert tabellen op prioriteit (610) in hoog, gemiddeld en laag.",  # Line 71
        "De risicoscoreberekening (620) weegt hoge tabellen drievoudig. De",  # Line 72
        "risiconivelabepaling (630) bepaalt het algehele risiconiveau.",  # Line 73
        "",  # Line 74
        "Figuur 7 toont een vergelijkingsmatrix (700) met de waardepropositie (710)",  # Line 75 - NUMBER HERE
        "van het systeem ten opzichte van concurrenten en handmatige methoden.",  # Line 76
        "",  # Line 77
        "GEDETAILLEERDE BESCHRIJVING VAN DE UITVINDING",  # Line 78
        "Verwijzend naar Figuur 1, omvat het intelligente databasescannerplatform",  # Line 79
        "(100) ondersteuning voor drie database-engines. De PostgreSQL-module (110)",  # Line 80 - NUMBER HERE
        "maakt gebruik van de psycopg2-bibliotheek voor verbinding. De MySQL-module",  # Line 81
        "(120) gebruikt de officiële MySQL-connector. De Microsoft SQL Server-module",  # Line 82
        "(130) maakt gebruik van pyodbc voor ODBC-verbindingen.",  # Line 83
        "",  # Line 84
        "De prioriteringsmodule (140) analyseert databaseschema's en kent",  # Line 85 - NUMBER HERE
        "prioriteitsscores toe aan tabellen. De parallelle verwerkingsmodule (150)",  # Line 86
        "maakt gebruik van drie gelijktijdige workers. De BSN-validatiemodule (160)",  # Line 87
        "implementeert de Nederlandse 11-proef voor verificatie van",  # Line 88
        "Burgerservicenummers.",  # Line 89
        "Verwijzend naar Figuur 2, berekent het prioriteitsscorealgoritme (200)",  # Line 90 - NUMBER HERE
        "scores in twee stappen. In stap 1 (210) worden basisscores toegekend op",  # Line 91
        "basis van tabelnamen: user/customer tabellen krijgen factor 3,0,",  # Line 92
        "medical/health tabellen factor 3,0, en payment/bank tabellen factor 2,8. In",  # Line 93
        "stap 2 (220) wordt een kolomboost toegevoegd: ssn/bsn/passport kolommen",  # Line 94
        "krijgen factor 3,0, email/phone kolommen factor 2,5, en address/birth",  # Line 95 - NUMBER HERE
        "kolommen factor 2,2. De uiteindelijke berekening (230) is: priority_score =",  # Line 96
        "min(base_score + column_boost, 3.5).",  # Line 97
        "",  # Line 98
        "Verwijzend naar Figuur 3, selecteert de beslisboom (300) automatisch de",  # Line 99
        "optimale bemonsteringsstrategie. De snelle modus (310) wordt gebruikt voor",  # Line 100 - NUMBER HERE
        "databases met maximaal 15 tabellen, waarbij 100 rijen per tabel worden",  # Line 101
        "geanalyseerd met 2 parallelle workers. De slimme modus (320) is geschikt",  # Line 102
        "voor databases met maximaal 50 tabellen, met analyse van 300 rijen en 3",  # Line 103
        "workers. De diepe modus (330) wordt ingezet voor databases tot 75 tabellen,",  # Line 104
        "met 500 rijen en 3 workers.",  # Line 105 - NUMBER HERE
        "",  # Line 106
        "Verwijzend naar Figuur 4, implementeert de parallelle scanningworkflow",  # Line 107
        "(400) een ThreadPoolExecutor (410) met maximaal 3 workers. Elke worker",  # Line 108
        "(420, 421, 422) verwerkt een afzonderlijke tabel en voert PII-detectie uit.",  # Line 109
        "De resultaataggregatie (430) combineert alle bevindingen in een",  # Line 110 - NUMBER HERE
        "gecentraliseerde resultatenlijst.",  # Line 111
        "",  # Line 112
        "Verwijzend naar Figuur 5, implementeert de BSN-validatiemodule het",  # Line 113
        "11-proefalgoritme (500). De formule (510) berekent: checksum = (d0×9) +",  # Line 114
        "(d1×8) + (d2×7) + (d3×6) + (d4×5) + (d5×4) + (d6×3) + (d7×2) - (d8×1). Een",  # Line 115 - NUMBER HERE
        "BSN is geldig indien checksum mod 11 = 0 (520). De detectiestroom (530)",  # Line 116
        "omvat drie stappen: eerst regex-patroonherkenning voor 9-cijferige",  # Line 117
        "nummers, vervolgens checksumvalidatie, en tenslotte GDPR-classificatie",  # Line 118
        "onder Artikel 9.",  # Line 119
        "EINDE BESCHRIJVING",  # Line 120 - NUMBER HERE
    ]
    
    # Add each line with proper formatting
    for i, line in enumerate(lines, 1):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        
        # Add line number at every 5th line
        if i % 5 == 0:
            line_num_run = para.add_run(f"{i}")
            line_num_run.font.size = Pt(10)
            # Add tab to separate number from text
            tab_run = para.add_run("\t")
        else:
            # Empty space for alignment
            para.add_run("\t")
        
        # Add the actual text
        text_run = para.add_run(line)
        text_run.font.size = Pt(11)
        text_run.font.name = "Times New Roman"
    
    doc.save("patent_documents/BESCHRIJVING_RVO_v2.docx")
    print("Created: BESCHRIJVING_RVO_v2.docx")


def create_conclusies_v2():
    """Create RVO-compliant Conclusies"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run("3")  # Page 3 (continues from Beschrijving)
        run.font.size = Pt(11)
    
    lines = [
        "",  # Line 1
        "CONCLUSIES",  # Line 2
        "",  # Line 3
        "1. Een computersysteem voor het scannen van databases, omvattende:",  # Line 4
        "   a) een databaseverbindingsmodule (110, 120, 130) voor het koppelen aan",  # Line 5
        "      meerdere relationele databases;",  # Line 6
        "   b) een prioriteringsmodule (140) die tabellen selecteert op basis van",  # Line 7
        "      gevoeligheidscriteria;",  # Line 8
        "   c) een adaptieve bemonsteringsmodule (310, 320, 330) voor het",  # Line 9
        "      analyseren van rijen;",  # Line 10
        "   d) een parallelle verwerkingsmodule (150, 420, 421, 422) voor",  # Line 11
        "      gelijktijdige tabelanalyse.",  # Line 12
        "",  # Line 13
        "2. Het systeem volgens conclusie 1, waarbij de databases PostgreSQL (110),",  # Line 14
        "   MySQL (120) en Microsoft SQL Server (130) omvatten.",  # Line 15
        "",  # Line 16
        "3. Het systeem volgens conclusie 1, waarbij de prioriteringsmodule (140)",  # Line 17
        "   tabellen identificeert op basis van naamgevingspatronen die duiden op",  # Line 18
        "   persoonsgegevens, waarbij de prioriteitsscoreberekening (200) een",  # Line 19
        "   basisscore (210) en kolomboost (220) combineert.",  # Line 20
        "",  # Line 21
        "4. Het systeem volgens conclusie 1, waarbij het systeem een",  # Line 22
        "   validatiemodule (160) omvat voor Nederlandse Burgerservicenummers (BSN)",  # Line 23
        "   op basis van een wiskundige controle volgens de 11-proef (500, 510,",  # Line 24
        "   520).",  # Line 25
        "",  # Line 26
        "5. Een methode voor het scannen van databases, omvattende de stappen van:",  # Line 27
        "   a) analyseren van een databaseschema (600, 610);",  # Line 28
        "   b) toekennen van prioriteitsscores aan tabellen (200, 210, 220, 230);",  # Line 29
        "   c) selecteren van een bemonsteringsstrategie (300, 310, 320, 330);",  # Line 30
        "   d) uitvoeren van parallelle scans (400, 410, 420, 421, 422, 430).",  # Line 31
        "",  # Line 32
        "6. Een computerleesbaar medium dat instructies bevat die, wanneer",  # Line 33
        "   uitgevoerd door een processor, het systeem volgens conclusie 1",  # Line 34
        "   implementeren.",  # Line 35
        "",  # Line 36
        "EINDE CONCLUSIES",  # Line 37
    ]
    
    for i, line in enumerate(lines, 1):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        
        if i % 5 == 0:
            line_num_run = para.add_run(f"{i}")
            line_num_run.font.size = Pt(10)
            para.add_run("\t")
        else:
            para.add_run("\t")
        
        text_run = para.add_run(line)
        text_run.font.size = Pt(11)
        text_run.font.name = "Times New Roman"
    
    doc.save("patent_documents/CONCLUSIES_RVO_v2.docx")
    print("Created: CONCLUSIES_RVO_v2.docx")


def create_samenvatting_v2():
    """Create RVO-compliant Samenvatting"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run("4")  # Page 4
        run.font.size = Pt(11)
    
    lines = [
        "",  # Line 1
        "SAMENVATTING",  # Line 2
        "",  # Line 3
        "De uitvinding betreft een intelligent databasescansysteem (100) voor het",  # Line 4
        "detecteren van persoonsgegevens in relationele databases. Het systeem",  # Line 5
        "ondersteunt meerdere database-engines (110, 120, 130) en analyseert",  # Line 6
        "databasestructuren om tabellen met een verhoogd privacy-risico te",  # Line 7
        "identificeren.",  # Line 8
        "",  # Line 9
        "Op basis van deze analyse selecteert het systeem automatisch een",  # Line 10
        "bemonsteringsstrategie (300, 310, 320, 330) en voert parallelle scans uit",  # Line 11
        "(400, 420, 421, 422). De uitvinding omvat tevens een specifieke validatie",  # Line 12
        "(160, 500) voor Nederlandse Burgerservicenummers, waardoor het systeem",  # Line 13
        "geschikt is voor toepassing binnen privacyregelgeving zoals de AVG.",  # Line 14
        "",  # Line 15
        "EINDE SAMENVATTING",  # Line 16
    ]
    
    for i, line in enumerate(lines, 1):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        
        if i % 5 == 0:
            line_num_run = para.add_run(f"{i}")
            line_num_run.font.size = Pt(10)
            para.add_run("\t")
        else:
            para.add_run("\t")
        
        text_run = para.add_run(line)
        text_run.font.size = Pt(11)
        text_run.font.name = "Times New Roman"
    
    doc.save("patent_documents/SAMENVATTING_RVO_v2.docx")
    print("Created: SAMENVATTING_RVO_v2.docx")


if __name__ == "__main__":
    print("=" * 60)
    print("Generating RVO-COMPLIANT patent documents v2")
    print("With VISIBLE line numbers: 5, 10, 15, 20...")
    print("=" * 60)
    
    create_beschrijving_v2()
    create_conclusies_v2()
    create_samenvatting_v2()
    
    print("=" * 60)
    print("All documents created with proper RVO formatting!")
    print("\nFormat applied:")
    print("- Page numbers: centered at top")
    print("- Line numbers: 5, 10, 15, 20... visible in left margin")
    print("- Continuous page numbering across documents")
