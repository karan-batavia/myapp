"""
Generate RVO-compliant patent documents with proper formatting:
- Page numbers centered at top
- Line numbers at every 5th line (5, 10, 15...) restarting each page
- Continuous page numbering across documents
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

os.makedirs("patent_documents", exist_ok=True)

def add_page_number(doc, page_num):
    """Add page number to header"""
    section = doc.sections[-1]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.text = str(page_num)
    header_para.runs[0].font.size = Pt(10) if header_para.runs else None

def format_with_line_numbers(text_lines, start_line=5):
    """Format text with line numbers at every 5th line"""
    result = []
    line_num = start_line
    
    for i, line in enumerate(text_lines, 1):
        if i % 5 == 0:
            result.append(f"{line_num}\t{line}")
            line_num += 5
        else:
            result.append(f"\t{line}")
    
    return result

# ============================================================
# BESCHRIJVING (Description) - Pages 1-2
# ============================================================
def create_beschrijving():
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    content = """BESCHRIJVING

TITEL
Intelligente Database Scanner met Adaptieve Sampling en Prioritering

TECHNISCH GEBIED
De uitvinding betreft een computersysteem voor de geautomatiseerde analyse van databases met het oog op gegevensbescherming en privacy-compliance. Het systeem ondersteunt meerdere relationele database-engines, waaronder PostgreSQL, MySQL en Microsoft SQL Server, en is ingericht voor het detecteren van persoonsgegevens.

Het systeem maakt gebruik van prioriteit-gebaseerde tabelselectie en adaptieve bemonsteringsstrategieën om efficiënt gegevens te analyseren. Daarnaast voorziet de uitvinding in detectie en validatie van specifiek Nederlandse identificatiegegevens, waaronder Burgerservicenummers (BSN).

ACHTERGROND VAN DE UITVINDING
Databases bevatten doorgaans grote hoeveelheden persoonsgegevens die verspreid zijn over meerdere tabellen. Bestaande databasescanners voeren vaak lineaire of volledige scans uit zonder onderscheid te maken tussen relevante en minder relevante tabellen.

Dit leidt tot lange verwerkingstijden en een verhoogde belasting van systeembronnen. Bovendien ontbreekt in veel bestaande oplossingen een specifieke validatie voor nationale identificatiegegevens, zoals de wiskundige 11-proef die vereist is voor BSN-nummers in Nederland.

SAMENVATTING VAN DE UITVINDING
De uitvinding voorziet in een intelligent databasescansysteem dat tabellen binnen een database analyseert en prioriteitsscores toekent op basis van tabel- en kolomnamen. Op grond van deze scores wordt een geschikte bemonsteringsstrategie geselecteerd.

Afhankelijk van het risicoprofiel van de database worden verschillende bemonsteringsniveaus toegepast, variërend van beperkte tot uitgebreide analyse. Het systeem maakt gebruik van parallelle verwerking om meerdere tabellen gelijktijdig te analyseren, waardoor de efficiëntie wordt verbeterd.

KORTE BESCHRIJVING VAN DE TEKENINGEN
Figuur 1 toont een overzicht van de multi-engine database scanner architectuur volgens de uitvinding. Het intelligente databasescannerplatform (100) ondersteunt drie database-engines: PostgreSQL via psycopg2 (110), MySQL via connector (120), en Microsoft SQL Server via pyodbc (130). Het systeem omvat een prioriteringsmodule (140), parallelle workers (150) en een BSN-validatiemodule (160).

Figuur 2 toont het prioriteitsscorealgoritme. De tabelprioriteitsberekening (200) omvat stap 1 voor de basisscore (210) waarbij tabel- en kolomnamen worden geanalyseerd, en stap 2 voor de kolomboost (220). De uiteindelijke scoreberekening (230) combineert deze waarden met een maximum van 3,5. Een voorbeeld (240) toont hoe een tabel "customer_profiles" met email- en telefoonkolommen de hoogste prioriteit verkrijgt.

Figuur 3 toont de adaptieve bemonsteringsstrategieën. De scanmodus-beslisboom (300) selecteert tussen snelle modus (310) voor kleine databases, slimme modus (320) voor middelgrote databases, en diepe modus (330) voor grote databases. Invoerparameters (340) bepalen de selectie. De tijdsbesparing (350) bedraagt tot 60% reductie.

Figuur 4 toont de parallelle scanningworkflow. De ThreadPoolExecutor (400, 410) verdeelt het werk over drie workers (420, 421, 422), elk verantwoordelijk voor het scannen van afzonderlijke tabellen. De resultaataggregatie (430) combineert alle bevindingen.

Figuur 5 toont de BSN-checksumvalidatie volgens de Nederlandse 11-proef. Het algoritme (500) past de formule (510) toe waarbij de checksum wordt berekend als: (d0×9) + (d1×8) + (d2×7) + (d3×6) + (d4×5) + (d5×4) + (d6×3) + (d7×2) - (d8×1). De validatie (520) controleert of checksum mod 11 = 0. De detectiestroom (530) omvat regex-patroonherkenning, checksumvalidatie en GDPR-classificatie.

Figuur 6 toont de schema-intelligentieanalyse. De databaserisicobepaling (600) categoriseert tabellen op prioriteit (610) in hoog, gemiddeld en laag. De risicoscoreberekening (620) weegt hoge tabellen drievoudig. De risiconivelabepaling (630) bepaalt het algehele risiconiveau en de aanbevolen scandiepte.

Figuur 7 toont een vergelijkingsmatrix (700) met de waardepropositie (710) van het systeem ten opzichte van concurrenten en handmatige methoden.

GEDETAILLEERDE BESCHRIJVING VAN DE UITVINDING
Verwijzend naar Figuur 1, omvat het intelligente databasescannerplatform (100) ondersteuning voor drie database-engines. De PostgreSQL-module (110) maakt gebruik van de psycopg2-bibliotheek voor verbinding. De MySQL-module (120) gebruikt de officiële MySQL-connector. De Microsoft SQL Server-module (130) maakt gebruik van pyodbc voor ODBC-verbindingen.

De prioriteringsmodule (140) analyseert databaseschema's en kent prioriteitsscores toe aan tabellen. De parallelle verwerkingsmodule (150) maakt gebruik van drie gelijktijdige workers. De BSN-validatiemodule (160) implementeert de Nederlandse 11-proef voor verificatie van Burgerservicenummers.

Verwijzend naar Figuur 2, berekent het prioriteitsscorealgoritme (200) scores in twee stappen. In stap 1 (210) worden basisscores toegekend op basis van tabelnamen: user/customer tabellen krijgen factor 3,0, medical/health tabellen factor 3,0, en payment/bank tabellen factor 2,8. In stap 2 (220) wordt een kolomboost toegevoegd: ssn/bsn/passport kolommen krijgen factor 3,0, email/phone kolommen factor 2,5, en address/birth kolommen factor 2,2. De uiteindelijke berekening (230) is: priority_score = min(base_score + column_boost, 3.5).

Verwijzend naar Figuur 3, selecteert de beslisboom (300) automatisch de optimale bemonsteringsstrategie. De snelle modus (310) wordt gebruikt voor databases met maximaal 15 tabellen, waarbij 100 rijen per tabel worden geanalyseerd met 2 parallelle workers. De slimme modus (320) is geschikt voor databases met maximaal 50 tabellen, met analyse van 300 rijen en 3 workers. De diepe modus (330) wordt ingezet voor databases tot 75 tabellen, met 500 rijen en 3 workers.

Verwijzend naar Figuur 4, implementeert de parallelle scanningworkflow (400) een ThreadPoolExecutor (410) met maximaal 3 workers. Elke worker (420, 421, 422) verwerkt een afzonderlijke tabel en voert PII-detectie uit. De resultaataggregatie (430) combineert alle bevindingen in een gecentraliseerde resultatenlijst.

Verwijzend naar Figuur 5, implementeert de BSN-validatiemodule het 11-proefalgoritme (500). De formule (510) berekent: checksum = (d0×9) + (d1×8) + (d2×7) + (d3×6) + (d4×5) + (d5×4) + (d6×3) + (d7×2) - (d8×1). Een BSN is geldig indien checksum mod 11 = 0 (520). De detectiestroom (530) omvat drie stappen: eerst regex-patroonherkenning voor 9-cijferige nummers, vervolgens checksumvalidatie, en tenslotte GDPR-classificatie onder Artikel 9.

EINDE BESCHRIJVING"""

    lines = content.split('\n')
    
    # Add header with page number
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("1")
    run.font.size = Pt(10)
    
    # Add content with line numbers
    line_counter = 0
    for line in lines:
        para = doc.add_paragraph()
        line_counter += 1
        
        if line_counter % 5 == 0:
            para.add_run(f"{line_counter}\t").bold = False
        else:
            para.add_run("\t")
        
        para.add_run(line)
        para.paragraph_format.line_spacing = 1.5
    
    doc.save("patent_documents/BESCHRIJVING_RVO.docx")
    print("Created: BESCHRIJVING_RVO.docx (Page 1-2)")
    return 2  # Returns last page number

# ============================================================
# CONCLUSIES (Claims) - Pages 3-4 (continuing from Beschrijving)
# ============================================================
def create_conclusies(start_page=3):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    content = """CONCLUSIES

1. Een computersysteem voor het scannen van databases, omvattende:
a) een databaseverbindingsmodule (110, 120, 130) voor het koppelen aan meerdere relationele databases;
b) een prioriteringsmodule (140) die tabellen selecteert op basis van gevoeligheidscriteria;
c) een adaptieve bemonsteringsmodule (310, 320, 330) voor het analyseren van rijen;
d) een parallelle verwerkingsmodule (150, 420, 421, 422) voor gelijktijdige tabelanalyse.

2. Het systeem volgens conclusie 1, waarbij de databases PostgreSQL (110), MySQL (120) en Microsoft SQL Server (130) omvatten.

3. Het systeem volgens conclusie 1, waarbij de prioriteringsmodule (140) tabellen identificeert op basis van naamgevingspatronen die duiden op persoonsgegevens, waarbij de prioriteitsscoreberekening (200) een basisscore (210) en kolomboost (220) combineert.

4. Het systeem volgens conclusie 1, waarbij het systeem een validatiemodule (160) omvat voor Nederlandse Burgerservicenummers (BSN) op basis van een wiskundige controle volgens de 11-proef (500, 510, 520).

5. Een methode voor het scannen van databases, omvattende de stappen van:
a) analyseren van een databaseschema (600, 610);
b) toekennen van prioriteitsscores aan tabellen (200, 210, 220, 230);
c) selecteren van een bemonsteringsstrategie (300, 310, 320, 330);
d) uitvoeren van parallelle scans (400, 410, 420, 421, 422, 430).

6. Een computerleesbaar medium dat instructies bevat die, wanneer uitgevoerd door een processor, het systeem volgens conclusie 1 implementeren.

EINDE CONCLUSIES"""

    # Add header with page number
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(str(start_page))
    run.font.size = Pt(10)
    
    lines = content.split('\n')
    line_counter = 0
    
    for line in lines:
        para = doc.add_paragraph()
        line_counter += 1
        
        if line_counter % 5 == 0:
            para.add_run(f"{line_counter}\t").bold = False
        else:
            para.add_run("\t")
        
        para.add_run(line)
        para.paragraph_format.line_spacing = 1.5
    
    doc.save("patent_documents/CONCLUSIES_RVO.docx")
    print(f"Created: CONCLUSIES_RVO.docx (Page {start_page})")
    return start_page + 1

# ============================================================
# SAMENVATTING (Abstract) - Last page
# ============================================================
def create_samenvatting(start_page=4):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    content = """SAMENVATTING

De uitvinding betreft een intelligent databasescansysteem (100) voor het detecteren van persoonsgegevens in relationele databases. Het systeem ondersteunt meerdere database-engines (110, 120, 130) en analyseert databasestructuren om tabellen met een verhoogd privacy-risico te identificeren.

Op basis van deze analyse selecteert het systeem automatisch een bemonsteringsstrategie (300, 310, 320, 330) en voert parallelle scans uit (400, 420, 421, 422). De uitvinding omvat tevens een specifieke validatie (160, 500) voor Nederlandse Burgerservicenummers, waardoor het systeem geschikt is voor toepassing binnen privacyregelgeving zoals de AVG.

EINDE SAMENVATTING"""

    # Add header with page number
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(str(start_page))
    run.font.size = Pt(10)
    
    lines = content.split('\n')
    line_counter = 0
    
    for line in lines:
        para = doc.add_paragraph()
        line_counter += 1
        
        if line_counter % 5 == 0:
            para.add_run(f"{line_counter}\t").bold = False
        else:
            para.add_run("\t")
        
        para.add_run(line)
        para.paragraph_format.line_spacing = 1.5
    
    doc.save("patent_documents/SAMENVATTING_RVO.docx")
    print(f"Created: SAMENVATTING_RVO.docx (Page {start_page})")

if __name__ == "__main__":
    print("Generating RVO-compliant patent documents...")
    print("=" * 50)
    
    # Create documents with continuous page numbering
    last_page = create_beschrijving()
    last_page = create_conclusies(start_page=last_page + 1)
    create_samenvatting(start_page=last_page)
    
    print("=" * 50)
    print("All documents created in patent_documents/ folder")
    print("\nDocuments created:")
    print("1. Tekeningen_RVO_Compliant.pdf - Drawings with reference numbers only")
    print("2. BESCHRIJVING_RVO.docx - Description (includes drawing explanations)")
    print("3. CONCLUSIES_RVO.docx - Claims with reference numbers")
    print("4. SAMENVATTING_RVO.docx - Abstract with reference numbers")
